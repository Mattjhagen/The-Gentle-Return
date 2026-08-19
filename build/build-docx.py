#!/usr/bin/env python3
"""
Build Amazon KDP-Compliant Word Manuscript (.docx) for The Gentle Return
Trim: 6" x 9", Mirrored Margins, Proper Styles, Front & Back Matter
"""

import os
import re
from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

PROJECT_DIR = Path(__file__).parent.parent
MANUSCRIPT_DIR = PROJECT_DIR / 'manuscript'
CHAPTERS_DIR = MANUSCRIPT_DIR / 'chapters'
OUTPUT_DIR = PROJECT_DIR / 'build' / 'output'
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def add_page_number(run):
    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def build_docx():
    docx_path = OUTPUT_DIR / 'the-gentle-return-kdp.docx'
    print(f'Building KDP 6x9 Word Document: {docx_path}')

    doc = docx.Document()

    # Configure 6x9 page setup on default section
    section = doc.sections[0]
    section.page_width = Inches(6.0)
    section.page_height = Inches(9.0)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)    # Gutter
    section.right_margin = Inches(0.625)  # Outside
    section.different_first_page_header_footer = True
    section.odd_and_even_pages_header_footer = True

    # Configure headers & footers
    header_odd = section.header
    p_odd = header_odd.paragraphs[0]
    p_odd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_odd = p_odd.add_run('THE GENTLE RETURN')
    r_odd.font.name = 'Palatino Linotype'
    r_odd.font.size = Pt(8.5)
    r_odd.font.italic = True
    r_odd.font.color.rgb = RGBColor(90, 90, 90)

    header_even = section.even_page_header
    p_even = header_even.paragraphs[0]
    p_even.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_even = p_even.add_run('MATTHEW JAMES HAGEN')
    r_even.font.name = 'Palatino Linotype'
    r_even.font.size = Pt(8.5)
    r_even.font.italic = True
    r_even.font.color.rgb = RGBColor(90, 90, 90)

    footer = section.footer
    p_foot = footer.paragraphs[0]
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_foot = p_foot.add_run()
    r_foot.font.name = 'Palatino Linotype'
    r_foot.font.size = Pt(9.5)
    add_page_number(r_foot)

    footer_even = section.even_page_footer
    p_foot_even = footer_even.paragraphs[0]
    p_foot_even.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_foot_even = p_foot_even.add_run()
    r_foot_even.font.name = 'Palatino Linotype'
    r_foot_even.font.size = Pt(9.5)
    add_page_number(r_foot_even)

    # Styles
    styles = doc.styles

    # Normal Style
    normal_style = styles['Normal']
    normal_style.font.name = 'Palatino Linotype'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(20, 20, 20)
    normal_style.paragraph_format.line_spacing = 1.3
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(0)
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def add_styled_paragraph(doc, text, first_line_indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if first_line_indent:
            p.paragraph_format.first_line_indent = Inches(0.25)
        else:
            p.paragraph_format.first_line_indent = Inches(0)

        # Parse inline bold / italics
        tokens = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|___.*?___|__.*?__|_[^_]+_)', text)
        for token in tokens:
            if not token:
                continue
            if token.startswith('***') and token.endswith('***'):
                r = p.add_run(token[3:-3])
                r.bold = True
                r.italic = True
            elif token.startswith('**') and token.endswith('**'):
                r = p.add_run(token[2:-2])
                r.bold = True
            elif token.startswith('*') and token.endswith('*'):
                r = p.add_run(token[1:-1])
                r.italic = True
            elif token.startswith('___') and token.endswith('___'):
                r = p.add_run(token[3:-3])
                r.bold = True
                r.italic = True
            elif token.startswith('__') and token.endswith('__'):
                r = p.add_run(token[2:-2])
                r.bold = True
            elif token.startswith('_') and token.endswith('_'):
                r = p.add_run(token[1:-1])
                r.italic = True
            else:
                clean_text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', lambda m: m.group(1), token)
                p.add_run(clean_text)
        return p

    def parse_md_to_docx(doc, md_text):
        lines = md_text.strip().splitlines()
        i = 0
        first_p = False

        while i < len(lines):
            line = lines[i].rstrip()
            stripped = line.strip()

            if not stripped:
                i += 1
                continue

            if stripped == '---':
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)
                r = p.add_run('*   *   *')
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(100, 100, 100)
                first_p = True
                i += 1
                continue

            if stripped.startswith('# '):
                h_text = stripped[2:].strip()
                p = doc.add_paragraph()
                p.paragraph_format.page_break_before = True
                p.paragraph_format.space_before = Pt(72)
                p.paragraph_format.space_after = Pt(24)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                if ':' in h_text and ('Chapter' in h_text or 'Prologue' in h_text or 'Epilogue' in h_text):
                    parts = h_text.split(':', 1)
                    num_run = p.add_run(parts[0].strip().upper() + "\n")
                    num_run.font.size = Pt(11)
                    num_run.font.color.rgb = RGBColor(100, 100, 100)
                    title_run = p.add_run(parts[1].strip())
                    title_run.font.size = Pt(18)
                    title_run.bold = True
                else:
                    r = p.add_run(h_text)
                    r.font.size = Pt(17)
                    r.bold = True
                first_p = True
                i += 1
                continue

            if stripped.startswith('## '):
                h2_text = stripped[3:].strip()
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(16)
                p.paragraph_format.space_after = Pt(8)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(h2_text)
                r.font.size = Pt(12)
                r.bold = True
                first_p = True
                i += 1
                continue

            if stripped.startswith('### '):
                h3_text = stripped[4:].strip()
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(h3_text)
                r.font.size = Pt(11)
                r.bold = True
                first_p = True
                i += 1
                continue

            if stripped.startswith('- ') or stripped.startswith('* '):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.add_run('• ' + stripped[2:].strip())
                i += 1
                continue

            if stripped.startswith('> '):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.left_indent = Inches(0.35)
                p.paragraph_format.right_indent = Inches(0.35)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                r = p.add_run(stripped[2:].strip())
                r.italic = True
                i += 1
                continue

            # Accumulate paragraph
            para_lines = [stripped]
            while i + 1 < len(lines):
                next_line = lines[i + 1].rstrip()
                next_stripped = next_line.strip()
                if not next_stripped or next_stripped.startswith('#') or next_stripped == '---' or next_stripped.startswith('- ') or next_stripped.startswith('> ') or re.match(r'^\d+\.\s+', next_stripped):
                    break
                para_lines.append(next_stripped)
                i += 1

            para_text = ' '.join(para_lines)
            if first_p:
                add_styled_paragraph(doc, para_text, first_line_indent=False)
                first_p = False
            else:
                add_styled_paragraph(doc, para_text, first_line_indent=True)

            i += 1

    # Add Front Matter
    front_file = MANUSCRIPT_DIR / 'front-matter.md'
    if front_file.exists():
        parse_md_to_docx(doc, front_file.read_text(encoding='utf-8'))

    # Add Chapters
    for ch_path in sorted(CHAPTERS_DIR.glob('*.md')):
        parse_md_to_docx(doc, ch_path.read_text(encoding='utf-8'))

    # Add Back Matter
    back_file = MANUSCRIPT_DIR / 'back-matter.md'
    if back_file.exists():
        parse_md_to_docx(doc, back_file.read_text(encoding='utf-8'))

    doc.save(str(docx_path))
    file_size_mb = docx_path.stat().st_size / (1024 * 1024)
    print(f'Successfully generated: {docx_path} ({file_size_mb:.2f} MB)')


if __name__ == '__main__':
    build_docx()
